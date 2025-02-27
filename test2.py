import streamlit as st
from pinecone import Pinecone
import os
import pdfplumber
import docx
import pydub
import speech_recognition as sr
from io import BytesIO
from sentence_transformers import SentenceTransformer
from transformers import pipeline
from ibm_watson import SpeechToTextV1
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
from dotenv import load_dotenv
import requests
import re

# === Initialization ===

#load environment variables
load_dotenv(dotenv_path="C:\\codes\\testStreamlit\\Key.env")
HUG_TOKEN = os.getenv("HUG_TOKEN")
IBM_API_KEY = os.getenv("IBM_API_KEY")
IBM_URL = os.getenv("IBM_URL")
pinecone_key = os.getenv("PINECONE_API_KEY")
pinecone_index = os.getenv("PINECONE_INDEX_NAME")
RAG_PATH = os.getenv("RAG_FILE")
ADMIN_PSW = os.getenv("ADMIN_PSW")
#init pydub
pydub.AudioSegment.ffmpeg = "C:/ffmpeg/bin/ffmpeg.exe"

#load embedding model
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

#hugging face API URL & headers
HF_API_URL = "https://api-inference.huggingface.co/models/meta-llama/Llama-3.2-3B-Instruct"
headers = {"Authorization": f"Bearer {HUG_TOKEN}"}

#init pinecone
pc = Pinecone(api_key=pinecone_key)
index = pc.Index(pinecone_index)

# Streamlit App Title
st.title("CIC Community Chatbot")

# IBM Watson Authentication
authenticator = IAMAuthenticator(IBM_API_KEY)
transcribe = SpeechToTextV1(authenticator=authenticator)
transcribe.set_service_url(IBM_URL)


# === Text Extraction Functions ===
def check_file_type(file_input):
    # If input is an UploadedFile object (from Streamlit)
    if hasattr(file_input, "name"):
        file_extension = os.path.splitext(file_input.name)[1].lower()
    else:
        file_extension = os.path.splitext(file_input)[1].lower()

    if file_extension == ".pdf":
        if hasattr(file_input, "read"):  # UploadedFile case
            with pdfplumber.open(file_input) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        else:  # File path case
            with pdfplumber.open(file_input) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

    elif file_extension in [".doc", ".docx"]:
        if hasattr(file_input, "read"):  # UploadedFile case
            doc = docx.Document(file_input)
        else:  # File path case
            doc = docx.Document(file_input)

        return "\n".join([para.text for para in doc.paragraphs])

    return "Unsupported file type."


#load RAG data
dummy_data = None
#rag_file = check_file_type(RAG_PATH)
#dummy_data = rag_file 


# === RAG Initialization ===
def name_chunks(data):
    summarizer = pipeline("summarization", model="t5-small")
    text = (data)
    output = summarizer(text, max_length=4, min_length=1, do_sample=False)
    return output[0]["summary_text"] 
    
def split_into_semantic_chunks(text, max_tokens=1300, overlap=200):
    """
    Splits text into chunks based on semantic boundaries (paragraphs)
    with a sliding window for overlapping context.
    """
    paragraphs = text.split("\n\n")  # Split by paragraphs
    chunks = []
    current_chunk = []

    for paragraph in paragraphs:
        tokens = embedding_model(paragraph, return_tensors='pt', truncation=False)['input_ids'][0]
        token_count = len(tokens)

        # Check if adding this paragraph exceeds the max token limit
        if sum(len(tokenizer(p, return_tensors='pt')['input_ids'][0]) for p in current_chunk) + token_count > max_tokens:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = [paragraph]
        else:
            current_chunk.append(paragraph)

    # Append the last chunk
    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    # Sliding window to create overlaps
    sliding_chunks = []
    for i in range(0, len(chunks), 1):  # Iterate by 1 to overlap all
        combined = " ".join(chunks[max(0, i-1):i+1])  # Combine current and previous
        sliding_chunks.append(combined)

    return sliding_chunks

def initialize_rag(data):
    """
    Initializes the RAG system by chunking the data, encoding it,
    and upserting it into the Pinecone index.
    """
    # Apply semantic chunking with sliding window
    chunks = split_into_semantic_chunks(data)
    
    # Encode chunks into embeddings
    embeddings = embedding_model.encode(chunks).tolist()
    
    # Name the chunks
    name = name_chunks(chunks)
    
    # Upsert chunks into Pinecone with order and name metadata
    for i, chunk in enumerate(chunks):
        index.upsert(vectors=[(f"{name} #{i}",embeddings[i], {"text": chunk})])
    return name, i+1

# Streamlit UI: Upsert Data
rag_data = None
psw = st.text_input("Enter Admin Password:", type="password")
upsert = st.button("Upsert Data");new_file = st.file_uploader("Change upsert path",type= '.pdf')

if new_file and upsert:
    dummy_data = check_file_type(new_file)
    initialize_rag(dummy_data)
    name, records = initialize_rag(dummy_data)
    if psw == ADMIN_PSW and records > 1:
        st.success(f"Upsert successful!You upserted {records} record of, {name} this name was AI generated btw, to your pinecone vector db")
    elif psw == ADMIN_PSW:
        st.success(f"Upsert successful!You upserted {records} records of, {name} this name was AI generated btw, to your pinecone vector db")
    elif psw != ADMIN_PSW:
        st.warning("Incorrect password. Access denied.")


# === Search Relevant Chunks ===
def search_relevant_chunks(query, top_k=5):
    #sanity checks
    if index is None:
        raise ValueError("Pinecone index is not initialized.")
    elif embedding_model is None:
        raise ValueError("Embedding model is not initialized.")

    #encode the query
    query_vector = embedding_model.encode([query]).tolist()[0]  # Ensure list format

    #query pinecone
    results = index.query(vector=query_vector, top_k=top_k, include_metadata=True)

    #extract relevant text chunks
    return [match["metadata"]["text"] for match in results["matches"]]


# === AI Response Generation ===
def get_ai_response(query, context):
    context = search_relevant_chunks(query)
    
    #prompt to pass to model
    full_prompt = f"""
    You are a helpful assistant. Answer the user's question in a neutral tone.
    
    IMPORTANT:
    - If no context is available , respond with exactly this: 
    "I don't have the required context to answer this."
    -If the retrieved context does not fit the question, give this exact disclaimer(after doing this try to answer the question):
    "I don't have enough context to answer this but im assuming:"
    - Do not include any other information, explanations, or context after the answer. 
    - Answer only the question based on the provided context.
    
    Context:
    {context}
    
    Question: {query}
    
    Answer:
    """

    #send request to hugging face API
    response = requests.post(
        HF_API_URL,
        headers=headers,
        json={
            "inputs": full_prompt,
            "parameters": {
                "max_new_tokens": 150,
                "temperature": 0.2,
            },
        },
    )

    if response.status_code == 200:
        #whole answer with context and question
        ai_response = response.json()[0]["generated_text"]
        print("Full AI Response:", ai_response)

        #extract the answer using re
        answer_match = re.search(r"Answer:\s*(.*)", ai_response, re.DOTALL)
        if answer_match:
            answer = answer_match.group(1).strip()
        else:
            answer_start = ai_response.find("Answer:")
            answer = ai_response[answer_start + len("Answer:") :].strip() if answer_start != -1 else "Error: Could not extract answer."

        #streamlit UI
        st.subheader("AI Response:")
        return answer
    else:
        st.error(f"Error generating response: {response.text}")
        return None


# === User Input Handling ===
st.header("Type a Question")

with st.form("question_form"):
    user_text = st.text_area("", "Ask a question and submit to get a response.")
    submitted = st.form_submit_button("Submit")

if submitted:
    ans = get_ai_response(user_text, rag_data)
    st.write(f"**Answer:** {ans}")


# === Audio Input Handling ===
st.header("Say a Question")
#gets audio from recorder or upload
user_input_type = st.radio("", ["📁 Upload", "🎙️ Record"])
audio = (
    st.file_uploader("Upload an audio file") if user_input_type == "📁 Upload" else st.audio_input("Record your question")
)
#store the audio
if audio:
    st.audio(audio)
    audio_bytes = BytesIO(audio.read())
    raw_audio = pydub.AudioSegment.from_file(audio_bytes, format="wav")
    wav_audio = BytesIO()
    raw_audio.export(wav_audio, format="wav")

    with sr.AudioFile(wav_audio) as source:
        recognizer = sr.Recognizer()
        speech_input = recognizer.record(source)
#transcribe audio then pass it to model
    try:
        wav_data = speech_input.get_wav_data()
        speech_to_txt = transcribe.recognize(audio=wav_data, content_type="audio/wav").get_result()

        if "results" in speech_to_txt and speech_to_txt["results"]:
            speech_text = speech_to_txt["results"][0]["alternatives"][0]["transcript"]
            st.write(f"Transcribed text: {speech_text}")

            if st.button("Generate Response"):
                aud_ans = get_ai_response(speech_text,rag_data)
                st.write(f'**Answer:{aud_ans}**')
        else:
            st.error("No transcript found.")

    except sr.UnknownValueError:
        st.error("Could not understand the audio.")
    except sr.RequestError as e:
        st.error(f"Request error: {e}")
